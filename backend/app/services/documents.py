from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.encryption import decrypt_bytes, decrypt_text, encrypt_bytes, encrypt_text
from app.db.models import DocumentType, StoredDocument, StoredDocumentVersion


def _detect_document_type(filename: str, mime_type: str | None) -> DocumentType:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or (mime_type and "pdf" in mime_type):
        return DocumentType.pdf
    if ext == ".docx" or (mime_type and "word" in mime_type):
        return DocumentType.docx
    if ext == ".txt" or (mime_type and "text/plain" in mime_type):
        return DocumentType.txt
    return DocumentType.other


def _extract_text(content: bytes, doc_type: DocumentType) -> str:
    if doc_type == DocumentType.txt:
        return content.decode("utf-8", errors="ignore")
    if doc_type == DocumentType.pdf:
        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if doc_type == DocumentType.docx:
        document = DocxDocument(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return ""


def create_or_update_document(
    db: Session,
    owner_type: str,
    owner_id: int,
    title: str,
    filename: str,
    content: bytes,
    mime_type: str,
    document_id: int | None = None,
) -> StoredDocument:
    settings = get_settings()
    storage_root = Path(settings.storage_root)
    storage_root.mkdir(parents=True, exist_ok=True)

    if document_id:
        document = db.get(StoredDocument, document_id)
        if not document:
            raise ValueError("Document not found")
        next_version = document.current_version + 1
    else:
        document = StoredDocument(
            owner_type=owner_type,
            owner_id=owner_id,
            title=title,
            document_type=_detect_document_type(filename, mime_type),
            current_version=1,
        )
        db.add(document)
        db.flush()
        next_version = 1

    doc_dir = storage_root / "documents" / str(document.id)
    doc_dir.mkdir(parents=True, exist_ok=True)

    encrypted_bytes = encrypt_bytes(content)
    target_path = doc_dir / f"v{next_version}.bin"
    target_path.write_bytes(encrypted_bytes)

    extracted_text = _extract_text(content, document.document_type)

    version = StoredDocumentVersion(
        document_id=document.id,
        version_number=next_version,
        original_filename=filename,
        mime_type=mime_type,
        encrypted_file_path=str(target_path),
        extracted_text_encrypted=encrypt_text(extracted_text),
        checksum_sha256=hashlib.sha256(content).hexdigest(),
        size_bytes=len(content),
    )
    db.add(version)
    document.current_version = next_version

    db.commit()
    db.refresh(document)
    return document


def get_document_text(db: Session, document_id: int, version_number: int | None = None) -> str:
    document = db.get(StoredDocument, document_id)
    if not document:
        raise ValueError("Document not found")

    version = None
    if version_number is None:
        for candidate in document.versions:
            if candidate.version_number == document.current_version:
                version = candidate
                break
    else:
        for candidate in document.versions:
            if candidate.version_number == version_number:
                version = candidate
                break

    if not version:
        raise ValueError("Version not found")

    return decrypt_text(version.extracted_text_encrypted)


def read_document_file(db: Session, document_id: int, version_number: int | None = None) -> bytes:
    document = db.get(StoredDocument, document_id)
    if not document:
        raise ValueError("Document not found")

    version = None
    target_version = version_number or document.current_version
    for candidate in document.versions:
        if candidate.version_number == target_version:
            version = candidate
            break

    if not version:
        raise ValueError("Version not found")

    encrypted_bytes = Path(version.encrypted_file_path).read_bytes()
    return decrypt_bytes(encrypted_bytes)
